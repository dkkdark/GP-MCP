from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain.embeddings import OpenAIEmbeddings
import os
from docx import Document as DocxDocument
from langchain.chat_models import ChatOpenAI
from dotenv import load_dotenv
import prompts

load_dotenv()

def load_docx_plain(filepath):
    doc = DocxDocument(filepath)
    full_text = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            # paragraph
            para = element.xpath(".//w:t")
            if para:
                text = ''.join([t.text for t in para if t.text])
                full_text.append(text.strip())
        elif element.tag.endswith('tbl'):
            # table
            table = element
            for row in table.xpath(".//w:tr"):
                cells = row.xpath(".//w:tc")
                row_text = []
                for cell in cells:
                    cell_text = ''.join([t.text for t in cell.xpath(".//w:t") if t.text])
                    row_text.append(cell_text.strip())
                full_text.append(' | '.join(row_text))

    return '\n'.join(full_text)

def extract_chunks():
    llm = ChatOpenAI(temperature=0, model="gpt-4.1-mini", api_key=os.getenv("OPENAI_API_KEY"))

    doc_text = load_docx_plain(".//data/tasks.docx")

    splitter = RecursiveCharacterTextSplitter(chunk_size=700, chunk_overlap=150)

    documents = [
        Document(page_content=doc_text, metadata={"source": "tasks.docx"})
    ]

    chunks = splitter.split_documents(documents)

    chunck_type_prompt = prompts.get_chunck_type_prompt()

    enriched_chunks = []
    for chunk in chunks:
        prompt = chunck_type_prompt.format(text=doc_text, chunk=chunk.page_content)
        chunk_type = llm.predict(prompt).strip().lower()
        chunk.metadata["type"] = chunk_type
        enriched_chunks.append(chunk)

    return enriched_chunks

def create_db(enriched_chunks):
    db_path = "./teaching_chroma_db"

    db = Chroma(persist_directory=db_path, embedding_function=OpenAIEmbeddings())

    def add_chunk_to_chroma(chunks):
        db.add_documents(
            texts=[chunk.page_content for chunk in chunks],
            metadatas=[chunk.metadata for chunk in chunks]
        )
    add_chunk_to_chroma(enriched_chunks)

    return db

def get_retriever(type):
    enriched_chunks = extract_chunks()
    db = create_db(enriched_chunks)
    
    retriever = db.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 5, "filter": {"type": type}}
    )
    return retriever


# Store each row as its own chunk + add table type



# prompt = build_extraction_prompt(doc_text, doc_id="doc_042")
# response = llm.invoke(prompt)

# json_result = parse_llm_json_response(response.content)

# chunks = prepare_for_vector_db(json_result)
# print("----------------CHUNKS----------------")
# print(chunks)

# def parse_llm_json_response(response_text):
#     try:
#         cleaned = response_text.strip().strip("```").strip("json").strip()
#         parsed = json.loads(cleaned)
#         return parsed
#     except json.JSONDecodeError as e:
#         print("JSON parsing failed:", e)
#         return None
    
# def prepare_for_vector_db(valid_chunks):
#     return [
#         {
#             "id": chunk["id"],
#             "content": chunk["text"],
#             "type": chunk["type"],
#             "doc_id": chunk["doc_id"],
#             "metadata": {
#                 **chunk["metadata"]
#             }
#         }
#         for chunk in valid_chunks
#     ]

# ------------------------------------------------------------

# embeddings = OpenAIEmbeddings()

# db_location = ".//chrome_langchain_db"
# add_documents = not os.path.exists(db_location)

# splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=300)

# raw_documents = SimpleDirectoryReader(
#     input_files=[".//data/tasks.docx"]
# ).load_data()

# documents = [
#     LangchainDocument(page_content=doc.get_content(), metadata={"source": "tasks.docx"})
#     for doc in raw_documents
# ]

# nodes = splitter.split_documents(documents)

# if add_documents:
#     db = Chroma.from_documents(nodes, embedding=embeddings, persist_directory=db_location)
# else:
#     db = Chroma(persist_directory=db_location, embedding_function=embeddings)
#     db.add_documents(nodes)

        
# retriever = db.as_retriever(
#     search_kwargs={"k": 5}
# )

